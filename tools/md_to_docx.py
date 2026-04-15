#!/usr/bin/env python3
"""Convert a Markdown report into a Word (.docx) document with an auto-updating TOC.

Usage:
  python tools/md_to_docx.py

This script reads docs/Final_Report.md and writes docs/Final_Report.docx.
It recognizes the marker line "[TOC]" to insert a TOC field.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def add_table_of_contents(doc: Document, depth: int = 3) -> None:
    """Insert a Word Table of Contents field (auto-updating)."""
    paragraph = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f"TOC \\o \"1-{depth}\" \\h \\z \\u")
    paragraph._p.addnext(fld)


def add_code_paragraph(doc: Document, text: str, style_name: str = "Code") -> None:
    p = doc.add_paragraph(style=style_name)
    p.add_run(text)


def ensure_code_style(doc: Document) -> None:
    if "Code" in doc.styles:
        return

    style = doc.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = "Courier New"
    font.size = Pt(9)



if __name__ == "__main__":
    import sys

    repo_root = Path(__file__).resolve().parents[1]
    md_file = repo_root / "docs" / "Final_Report.md"
    out_file = repo_root / "docs" / "Final_Report.docx"

    # Create a new document
    doc = Document()

    ensure_code_style(doc)

    in_code_block = False

    # Precompute patterns
    heading_re = re.compile(r"^(#{1,6})\s+(.*)")
    bullet_re = re.compile(r"^\s*([*\-])\s+(.*)")
    number_re = re.compile(r"^\s*(\d+)\.\s+(.*)")

    with md_file.open(encoding="utf-8") as f:
        for raw_line in f:
            line = raw_line.rstrip("\n")

            # Honor the TOC insertion marker
            if line.strip() == "[TOC]":
                add_table_of_contents(doc, depth=3)
                continue

            # Code block toggling
            if line.strip().startswith("```"):
                in_code_block = not in_code_block
                if in_code_block:
                    # Start code block; add a blank line as separator
                    doc.add_paragraph()
                continue

            if in_code_block:
                # Add each line as a code line
                add_code_paragraph(doc, line)
                continue

            # Headings
            m = heading_re.match(line)
            if m:
                hashes, text = m.groups()
                level = min(len(hashes), 9)
                doc.add_heading(text.strip(), level=level)
                continue

            # Bullet list items
            m = bullet_re.match(line)
            if m:
                _, text = m.groups()
                doc.add_paragraph(text.strip(), style="List Bullet")
                continue

            # Numbered list items
            m = number_re.match(line)
            if m:
                _, text = m.groups()
                doc.add_paragraph(text.strip(), style="List Number")
                continue

            # Blank lines -> paragraph break
            if line.strip() == "":
                doc.add_paragraph()
                continue

            # Regular paragraph
            doc.add_paragraph(line)

    doc.save(out_file)
    print(f"Created {out_file}")
